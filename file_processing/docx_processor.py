from file_processor_strategy import FileProcessorStrategy
from docx import Document
import pandas as pd
class DocxFileProcessor(FileProcessorStrategy):
    def __init__(self, file_path: str) -> None:
        super().__init__(file_path)
        self.metadata = {}

    def process(self) -> None:
        doc = Document(self.file_path)
        self.metadata.update({'text': self.extract_text_from_docx(doc)})
        self.metadata.update({'author': doc.core_properties.author})
        self.metadata.update({'last_modified_by': doc.core_properties.last_modified_by})

        # Other core properties to include: https://python-docx.readthedocs.io/en/latest/api/document.html#coreproperties-objects
        # keywords, language, subject, version

    @staticmethod
    def extract_text_from_docx(doc: Document) -> str:
        try:
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error encountered while opening or processing {file_path}: {e}")
            return None

    def extract_tbls_from_docx(doc: Document) -> str:
        try:
            writer = pd.ExcelWriter(file_path.rsplit('.', 1)[0].lower()+'_tbls.xlsx', engine='xlsxwriter', engine_kwargs={'options': {'encoding': 'utf8'}})
            # for each table in doc
            for count, table in enumerate(doc.tables):
                df = pd.DataFrame()
                # write rows of table to df
                for row in table.rows:
                    text = [cell.text for cell in row.cells]
                    df = pd.concat([df, pd.DataFrame([text])], ignore_index=True)
                # write df to excel file
                df.to_excel(writer, sheet_name=file_path.rsplit('.', 1)[0].lower()[:20]+"_tbl_"+str(count))
            writer.close()
        except Exception as e:
            print(f"Error encountered while opening or processing {file_path}: {e}")
            return None